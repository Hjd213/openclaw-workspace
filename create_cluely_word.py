from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# 标题
para = doc.add_paragraph()
para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = para.add_run('Cluely 竞品分析报告')
run.font.size = Pt(18)
run.font.bold = True
run.font.name = '微软雅黑'

# 作者信息
para = doc.add_paragraph()
para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = para.add_run('作者：黄健东    日期：2026 年 3 月 11 日    版本：1.0（深度体验版）')
run.font.size = Pt(12)

# 1. 产品概述
para = doc.add_paragraph()
run = para.add_run('1. 产品概述')
run.font.size = Pt(14)
run.font.bold = True

para = doc.add_paragraph('1.1 产品定位')
para.runs[0].font.bold = True
para.add_run('\nCluely 是一款实时 AI 会议助手，主打"不可检测"的会议录制和实时 AI 辅助。与其他会议助手不同，Cluely 不需要 Bot 加入会议，而是通过系统级音频捕获实现隐蔽录制。')

para = doc.add_paragraph('\n核心差异化：')
para.runs[0].font.bold = True
para.add_run('\n• 免费版：不可检测（无需 Bot 加入）')
para.add_run('\n• Pro+ 版：屏幕共享时也能隐藏 Cluely（Undetectability 模式）')

para = doc.add_paragraph('\n1.2 目标用户')
para.runs[0].font.bold = True
para.add_run('\n• 需要隐蔽记录会议的专业人士')
para.add_run('\n• 追求高效操作的职场人士')
para.add_run('\n• 需要实时 AI 辅助的场景（面试、谈判、重要会议）')

para = doc.add_paragraph('\n1.3 核心功能')
para.runs[0].font.bold = True
para.add_run('\n• 实时会议转录（无需 Bot 加入）')
para.add_run('\n• 实时 AI 问答（会议中即时提问）')
para.add_run('\n• 自动生成会议纪要和摘要')
para.add_run('\n• 丰富的快捷键操作')
para.add_run('\n• 一键复制转录文本')
para.add_run('\n• 场景模板（Sales/Recruiting/Team Meet/Lecture 等）')

para = doc.add_paragraph('\n1.4 定价信息')
para.runs[0].font.bold = True
para.add_run('\n免费版：$0 - 有限 AI 响应次数、有限免费会议')
para.add_run('\nPro 版：$12/月（年付） - 无限 AI 响应、无限会议、最新 AI 模型、优先客服支持')
para.add_run('\nPro + Undetectability：$42/月（年付） - Pro 全部功能 + 屏幕共享隐藏 Cluely')
para.add_run('\n\n说明：以上为年付价格，月付价格更高（Pro 原价$20/月，Pro+ 原价$75/月）。免费版无需信用卡，同样支持不可检测（但屏幕共享时可能暴露）。')

# 2. 注册与上手体验
para = doc.add_paragraph()
run = para.add_run('2. 注册与上手体验')
run.font.size = Pt(14)
run.font.bold = True

para = doc.add_paragraph('2.1 注册流程')
para.runs[0].font.bold = True
para.add_run('\n• 注册方式：邮箱 / Google / Apple')
para.add_run('\n• 是否需要信用卡：否')
para.add_run('\n• 注册耗时：约 2 分钟')

para = doc.add_paragraph('\n2.2 初次使用引导')
para.runs[0].font.bold = True
para.add_run('\n• Onboarding 教程：有基础引导')
para.add_run('\n• 上手难度：2/5 分 - 界面简洁，容易上手')

para = doc.add_paragraph('\n2.3 平台支持')
para.runs[0].font.bold = True
para.add_run('\n• Windows 桌面应用')
para.add_run('\n• Mac 桌面应用')
para.add_run('\n• iOS 手机版')
para.add_run('\n• Web 版（app.cluely.com）')
para.add_run('\n• Android：不支持')

# 3. 功能体验评价
para = doc.add_paragraph()
run = para.add_run('3. 功能体验评价')
run.font.size = Pt(14)
run.font.bold = True

para = doc.add_paragraph('3.1 好用的功能')
para.runs[0].font.bold = True

features_good = [
    ('【实时转录】评分：5/5', '准确率非常高，二倍速播放依然准确'),
    ('【低延迟】评分：5/5', '延迟仅 1-2 个单词，几乎实时'),
    ('【不可检测】评分：5/5', '无需 Bot 加入，其他人看不到（免费版也支持，Pro+ 支持屏幕共享隐藏）'),
    ('【快捷键丰富】评分：5/5', 'Ctrl+\\显示/隐藏，Ctrl+ 方向键调整窗口'),
    ('【转录中 Chat】评分：4/5', '支持中文提问和中文回答'),
    ('【一键复制】评分：4/5', '转录文本可快速复制'),
    ('【AI 摘要】评分：4/5', '自动生成会议总结'),
    ('【场景模板】评分：4/5', '提供 Sales/Recruiting/Team Meet/Lecture 等 5 个预设模板，可自定义'),
]

for title, desc in features_good:
    para = doc.add_paragraph()
    para.add_run(title + '\n').font.bold = True
    para.add_run(desc)

para = doc.add_paragraph('\n3.2 不足的功能')
para.runs[0].font.bold = True

features_bad = [
    ('【总结界面 Chat】评分：2/5', '中文提问只能英文回答（Bug）'),
    ('【Transcription Language】评分：1/5', '不支持中文转录'),
    ('【跨会议记忆】评分：1/5', 'AI 只能针对当前会议提问'),
    ('【付费版价格】评分：3/5', 'Pro $12/月，Pro+ $42/月（年付），价格偏高'),
    ('【企业安全信息】评分：1/5', '未公开安全认证和数据政策'),
]

for title, desc in features_bad:
    para = doc.add_paragraph()
    para.add_run(title + '\n').font.bold = True
    para.add_run(desc)

# 4. 竞品对比
para = doc.add_paragraph()
run = para.add_run('4. 竞品对比')
run.font.size = Pt(14)
run.font.bold = True

para = doc.add_paragraph('4.1 Cluely 优势')
para.runs[0].font.bold = True
para.add_run('\n• 不可检测（独家优势，免费版也支持）')
para.add_run('\n• 快捷键丰富')
para.add_run('\n• 实时 AI 问答，低延迟')
para.add_run('\n• 转录准确率高（二倍速依然准确）')
para.add_run('\n• 免费版无需信用卡')
para.add_run('\n• 场景模板丰富（5 个预设）')

para = doc.add_paragraph('\n4.2 Cluely 劣势')
para.runs[0].font.bold = True
para.add_run('\n• 总结界面中文 Bug')
para.add_run('\n• 不支持中文转录')
para.add_run('\n• 无跨会议记忆')
para.add_run('\n• 企业安全信息未公开')
para.add_run('\n• 付费版价格偏高')

# 5. 总结与建议
para = doc.add_paragraph()
run = para.add_run('5. 总结与建议')
run.font.size = Pt(14)
run.font.bold = True

para = doc.add_paragraph('5.1 推荐场景')
para.runs[0].font.bold = True
para.add_run('\n• 需要隐蔽记录会议（面试、谈判、敏感会议）')
para.add_run('\n• 追求高效操作（快捷键友好）')
para.add_run('\n• 需要实时 AI 辅助（会议中即时提问）')
para.add_run('\n• 英文会议场景')

para = doc.add_paragraph('\n5.2 不推荐场景')
para.runs[0].font.bold = True
para.add_run('\n• 中文会议（不支持中文转录）')
para.add_run('\n• 需要跨会议检索（无历史记忆）')
para.add_run('\n• 企业用户（安全信息未公开）')
para.add_run('\n• 需要视频录制（无此功能）')

para = doc.add_paragraph('\n5.3 性价比评价')
para.runs[0].font.bold = True
para.add_run('\n免费版：值得使用，无需信用卡，适合轻度用户')
para.add_run('\n付费版：$12-$42/月（年付），价格偏高，Pro+ 屏幕共享隐藏是核心卖点')

para = doc.add_paragraph('\n5.4 最终评分')
para.runs[0].font.bold = True
para.add_run('\n总体评分：4.5/5')

# 6. 关键发现
para = doc.add_paragraph()
run = para.add_run('6. 关键发现')
run.font.size = Pt(14)
run.font.bold = True

para = doc.add_paragraph('1. 二倍速转录依然准确')
para.add_run('\n2. 延迟约 1-2 个单词')
para.add_run('\n3. 总结界面存在中文 Bug')
para.add_run('\n4. 免费会议次数限制')
para.add_run('\n5. Pro+ Undetectability 是核心升级功能（屏幕共享时隐藏 Cluely）')
para.add_run('\n6. 场景模板 - 提供 Sales/Recruiting/Team Meet/Lecture 等 5 个预设模板')

# 7. 对导师项目的启发
para = doc.add_paragraph()
run = para.add_run('7. 对导师项目的启发')
run.font.size = Pt(14)
run.font.bold = True

para = doc.add_paragraph('1. 语言支持需要架构级设计，不是补丁式（Cluely 中文 Bug 教训）')
para.add_run('\n2. 低延迟（1-2 词）是实时 AI 辅助的关键指标')
para.add_run('\n3. 场景模板设计值得借鉴（医疗/金融/办公等场景）')
para.add_run('\n4. 不可检测是差异化优势（屏幕共享隐藏是 Pro+ 卖点）')

# 测试环境
para = doc.add_paragraph()
run = para.add_run('测试环境')
run.font.size = Pt(14)
run.font.bold = True

para = doc.add_paragraph('测试日期：2026-03-11')
para.add_run('\n测试内容：B 站英文视频（8 分 30 秒）')
para.add_run('\n测试项目：实时转录、AI 问答、延迟测试、二倍速测试、场景模板')

para = doc.add_paragraph('\n参考资料：https://cluely.com | https://app.cluely.com')
para.add_run('\n\n报告完成日期：2026-03-11    体验时长：约 2 小时    可信度：5 星（实际注册 + 深度体验）')

doc.save(r'C:\Users\bljd5\Desktop\产品分析\Cluely 竞品分析报告.docx')
print('完成！')
