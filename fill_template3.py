from docx import Document
import os
import shutil

template_path = r'C:\Users\bljd5\Desktop\简历\稻小壳求职意向.docx'
temp_output = r'C:\Users\bljd5\.openclaw\workspace\黄健东 - 简历（求职意向）_temp.docx'
final_output = r'C:\Users\bljd5\Desktop\简历\黄健东 - 简历（求职意向）.docx'

# 先删除可能存在的旧文件
if os.path.exists(final_output):
    try:
        os.remove(final_output)
    except:
        pass

doc = Document(template_path)

replacements = {
    '稻小壳': '黄健东',
    'kingsoft@.com': '120241280105@ncepu.edu.cn',
    '12345678901': '+86 19973510193',
    'XX 大学': '华北电力大学',
    'XX 专业': '计算机/网络安全',
    '20XX': '2024',
    'XX 有限公司': '清华大学',
    'XX 职位': '科研助理',
    '大学英语四级证书、全国计算机等级考试证书（C 语言）': '大学英语六级（CET-6）、大学英语四级（CET-4）',
    '20XX-20XX 学年': '2024-2025 学年',
    '获得国家励志奖学金': '获校级立项（录取率<15%）',
    '获得学校一等奖学金': 'GPA 3.95/5.0，专业排名第 3',
    '性格开朗，喜欢与人交流，适应能力强': '乐观积极、勤奋好学，对计算机技术有浓厚热情',
    '具备良好的沟通能力和表达能力，学习能力强': '具备良好的 C 语言编程能力、沟通表达能力和团队协作经验',
    '具有团队合作精神，能适应部门之间的合作，具有很强的领导能力和执行能力': '专业基础扎实，在数学建模、网络安全竞赛中有实际项目经验'
}

for para in doc.paragraphs:
    for old, new in replacements.items():
        if old in para.text:
            para.text = para.text.replace(old, new)

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for old, new in replacements.items():
                if old in cell.text:
                    cell.text = cell.text.replace(old, new)

doc.save(temp_output)
shutil.copy(temp_output, final_output)
print(f'已完成：{final_output}')
