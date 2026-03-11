from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# 设置页面边距
section = doc.sections[0]
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

# 姓名
para = doc.add_paragraph()
para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = para.add_run('黄健东')
run.font.size = Pt(22)
run.font.bold = True
run.font.name = '微软雅黑'

# 基本信息
para = doc.add_paragraph()
para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = para.add_run('求职意向：软件开发/网络安全工程师')
run.font.size = Pt(10.5)
run.font.name = '宋体'

para = doc.add_paragraph()
para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = para.add_run('手机：+86 19973510193    邮箱：120241280105@ncepu.edu.cn    地址：中国北京市海淀区')
run.font.size = Pt(10.5)
run.font.name = '宋体'

# 教育背景
para = doc.add_paragraph()
run = para.add_run('教育背景')
run.font.size = Pt(12)
run.font.bold = True
run.font.name = '微软雅黑'

para = doc.add_paragraph('华北电力大学 · 计算机/网络安全 · 工学学士    2024.09 - 至今')
para.runs[0].font.size = Pt(10.5)
para.runs[0].font.name = '宋体'

para = doc.add_paragraph('GPA：3.95/5.0，专业排名第 3')
para.runs[0].font.size = Pt(10.5)

para = doc.add_paragraph('主修课程：高等数学（98 分）、C 程序设计（92 分）、程序设计实验（92 分）、计算机导论（90 分）')
para.runs[0].font.size = Pt(10.5)

# 项目经历
para = doc.add_paragraph()
run = para.add_run('项目经历')
run.font.size = Pt(12)
run.font.bold = True
run.font.name = '微软雅黑'

para = doc.add_paragraph('清华大学 | 科研助理    2025.07 - 2025.08')
para.runs[0].font.size = Pt(10.5)
para.runs[0].font.bold = True

para = doc.add_paragraph('    • 主导中美日三国数智医疗政策对比分析，提炼可借鉴策略')
para.runs[0].font.size = Pt(10.5)

para = doc.add_paragraph('    • 系统分析 20+ 项关键政策，撰写逾万字政策分析报告及 2 份医疗 AI 国家战略分析报告')
para.runs[0].font.size = Pt(10.5)

para = doc.add_paragraph('    • 熟练运用 AI 工具辅助研究')
para.runs[0].font.size = Pt(10.5)

para = doc.add_paragraph('国家级大学生创新创业训练计划 | 项目负责人    2024.09 - 至今')
para.runs[0].font.size = Pt(10.5)
para.runs[0].font.bold = True

para = doc.add_paragraph('    • 获校级立项（录取率<15%），组建 4 人跨学科团队')
para.runs[0].font.size = Pt(10.5)

para = doc.add_paragraph('    • 实现非遗传播年轻化核心目标，完成数字化与交互程序设计')
para.runs[0].font.size = Pt(10.5)

para = doc.add_paragraph('全国网络安全 CTF 竞赛    2024.08 - 至今')
para.runs[0].font.size = Pt(10.5)
para.runs[0].font.bold = True

para = doc.add_paragraph('    • 掌握虚拟机配置、Kali Linux 系统操作，具备 C 语言编程基础')
para.runs[0].font.size = Pt(10.5)

para = doc.add_paragraph('中国大学生数学建模竞赛    2025.04 - 2025.09')
para.runs[0].font.size = Pt(10.5)
para.runs[0].font.bold = True

para = doc.add_paragraph('    • 72 小时限时竞赛，负责数学建模与算法开发，使用 MATLAB 完成模型求解')
para.runs[0].font.size = Pt(10.5)

# 技能证书
para = doc.add_paragraph()
run = para.add_run('技能与证书')
run.font.size = Pt(12)
run.font.bold = True
run.font.name = '微软雅黑'

para = doc.add_paragraph('编程语言：C、MATLAB')
para.runs[0].font.size = Pt(10.5)

para = doc.add_paragraph('开发工具：Visual Studio、C-Free、AI 工具')
para.runs[0].font.size = Pt(10.5)

para = doc.add_paragraph('证书：大学英语六级（CET-6）、大学英语四级（CET-4）')
para.runs[0].font.size = Pt(10.5)

# 自我评价
para = doc.add_paragraph()
run = para.add_run('自我评价')
run.font.size = Pt(12)
run.font.bold = True
run.font.name = '微软雅黑'

para = doc.add_paragraph('乐观积极、勤奋好学，成绩优异。对计算机技术有浓厚热情，专业基础扎实。具备良好的 C 语言编程能力、沟通表达能力和团队协作经验。')
para.runs[0].font.size = Pt(10.5)

doc.save(r'C:\Users\bljd5\Desktop\简历\黄健东 - 简历（求职意向）.docx')
print('完成！')
